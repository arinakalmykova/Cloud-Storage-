from __future__ import annotations

import os
import shutil
from pathlib import Path

from docx import Document


def append_citation(paragraph, citation: str) -> bool:
    text = paragraph.text.strip()
    if not text or citation in text:
        return False

    if paragraph.runs:
        paragraph.runs[-1].text = paragraph.runs[-1].text.rstrip() + f" {citation}"
    else:
        paragraph.add_run(citation)
    return True


def main() -> None:
    source = Path(os.environ["SRC_DOCX"])
    output = Path(os.environ["DST_DOCX"])

    shutil.copy2(source, output)
    doc = Document(output)

    targets = [
        ("В данной работе рассматриваются четыре основных формата: JPEG, PNG, WebP и AVIF.", "[4]"),
        ("В качестве серверной технологии выбран язык программирования PHP с использованием фреймворка Laravel.", "[7]"),
        ("При разрабатываемого сервиса выбрано объектное хранилище S3 с использованием MinIO", "[3]"),
        ("CompressionService взаимодействует с PhotoService асинхронно через Kafka", "[2]"),
        ("формирование JWT-токена", "[9]"),
        ("критериям визуального качества и уменьшения размера файла", "[5]"),
        ("Модуль main.py обеспечивает HTTP-интерфейс FastAPI", "[8]"),
        ("каждый микросервис имеет собственную базу данных", "[1]"),
        ("Все необходимые параметры передаются через сообщения Kafka, а результаты обработки отправляются в виде событий.", "[2]"),
        ("Классификация выполняется микросервисом MLService на основе модели MobileNetV2.", "[6, 10]"),
    ]

    used = set()
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        for phrase, citation in targets:
            if phrase in text and citation not in used:
                if append_citation(paragraph, citation):
                    used.add(citation)

    doc.save(output)
    print(f"saved: {output}")
    print("citations added:", ", ".join(sorted(used)))


if __name__ == "__main__":
    main()
