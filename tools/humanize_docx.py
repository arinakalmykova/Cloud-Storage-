from __future__ import annotations

import os
import re
import shutil
from pathlib import Path

from docx import Document


CAPTION_PREFIXES = ("Рисунок ", "Таблица ")
LIST_PREFIXES = ("•", "-", "1)", "2)", "3)", "4)", "5)", "6)", "7)", "8)", "9)", "10)")
SKIP_STYLES = {"Heading 1", "Heading 2", "Heading 3", "Heading 4", "List Paragraph", "No Spacing"}


def lower_first(value: str) -> str:
    return value[:1].lower() + value[1:] if value else value


def humanize(text: str, idx: int) -> str:
    result = text.strip()

    replacements: list[tuple[str, str]] = [
        (r"\bВ данной главе рассматривается\b", "В этой главе рассматривается"),
        (r"\bВ данной главе описываются\b", "В этой главе описываются"),
        (r"\bВ данной главе\b", "В этой главе"),
        (r"\bв данной работе\b", "в этой работе"),
        (r"\bВ рамках\b", "При"),
        (r"\bв рамках\b", "при"),
        (r"\bСледует отметить\b", "Важно"),
        (r"\bследует отметить\b", "важно"),
        (r"\bна сегодняшний день\b", "сейчас"),
        (r"\bТаким образом\b", "В итоге"),
        (r"\bтаким образом\b", "в итоге"),
        (r"\bДанный\b", "Этот"),
        (r"\bданный\b", "этот"),
        (r"\bданная\b", "эта"),
        (r"\bданное\b", "это"),
        (r"\bданные\b", "эти"),
        (r"\bосуществляется\b", "выполняется" if idx % 2 == 0 else "происходит"),
        (r"\bосуществляются\b", "выполняются" if idx % 2 == 0 else "происходят"),
        (r"\bпредставлена\b", "показана" if idx % 2 == 0 else "представлена"),
        (r"\bпредставлены\b", "показаны" if idx % 2 == 0 else "представлены"),
        (r"\bпредставлен\b", "показан" if idx % 2 == 0 else "представлен"),
        (r"\bиспользуется\b", "применяется" if idx % 2 == 0 else "используется"),
        (r"\bиспользуются\b", "применяются" if idx % 2 == 0 else "используются"),
    ]
    for pattern, replacement in replacements:
        result = re.sub(pattern, replacement, result)

    sentences = re.split(r"(?<=[.!?])\s+", result)
    sentences = [sentence for sentence in sentences if sentence]

    if len(sentences) >= 2 and len(result) > 120:
        if idx % 5 == 0 and not sentences[1].startswith(("Примечательно, что", "Что немаловажно,", "Важно, ")):
            sentences[1] = "Примечательно, что " + lower_first(sentences[1])
        elif idx % 5 == 1 and not sentences[1].startswith(("Что немаловажно,", "Примечательно, что", "Важно, ")):
            sentences[1] = "Что немаловажно, " + lower_first(sentences[1])

    if sentences and len(sentences[0]) > 190 and "," in sentences[0] and idx % 4 == 0:
        sentences[0] = sentences[0].replace(",", ". ", 1)

    if len(sentences) >= 3 and idx % 6 == 0:
        third_sentence = sentences[2]
        if not third_sentence.startswith(("Кстати,", "При этом", "Впрочем,")):
            sentences[2] = "При этом " + lower_first(third_sentence)

    result = " ".join(sentences)
    result = re.sub(r"\s{2,}", " ", result)
    result = result.replace(" .", ".").replace(" ,", ",")
    return result.strip()


def should_skip_paragraph(text: str, style_name: str) -> bool:
    if not text:
        return True
    if style_name in SKIP_STYLES:
        return True
    if text.startswith(CAPTION_PREFIXES) or text.startswith(LIST_PREFIXES):
        return True
    if text.isupper():
        return True
    if len(text) < 45:
        return True
    if "http://" in text or "https://" in text:
        return True
    return False


def rewrite_paragraph(paragraph, new_text: str) -> None:
    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def main() -> None:
    source = Path(os.environ["SRC_DOCX"])
    out = Path(os.environ["DST_DOCX"])

    shutil.copy2(source, out)
    document = Document(out)

    processed = 0
    skipped = 0

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        style_name = paragraph.style.name if paragraph.style else ""
        if should_skip_paragraph(text, style_name):
            skipped += 1
            continue

        new_text = humanize(text, processed)
        if new_text != text:
            rewrite_paragraph(paragraph, new_text)
            processed += 1

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text.strip()
                    style_name = paragraph.style.name if paragraph.style else ""
                    if not text or len(text) < 70 or style_name in SKIP_STYLES:
                        continue
                    if text.startswith(CAPTION_PREFIXES) or text.startswith(LIST_PREFIXES):
                        continue

                    new_text = humanize(text, processed)
                    if new_text != text:
                        rewrite_paragraph(paragraph, new_text)
                        processed += 1

    document.save(out)
    print(f"saved: {out}")
    print(f"processed paragraphs: {processed}")
    print(f"skipped paragraphs: {skipped}")


if __name__ == "__main__":
    main()
