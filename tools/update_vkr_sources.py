from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ACCESS_DATE = "18.02.2026"

SOURCES = [
    (
        "1) Ворсин В. А. Микро-сервисная архитектура бизнес-приложений: перспективы и проблемы "
        "[Электронный ресурс] // GLOBUS. 2020. С. 51-53. URL: "
        "https://cyberleninka.ru/article/n/mikro-servisnaya-arhitektura-biznes-prilozheniy-perspektivy-i-problemy "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "2) Карпович М. Н. Особенности проектирования микросервисно-событийных архитектур для "
        "высоконагруженных распределенных систем обработки информаций [Электронный ресурс] // "
        "Труды БГТУ. Серия 3: Физико-математические науки и информатика. 2023. № 1 (266). С. 89-95. "
        "DOI: 10.52065/2520-6141-2023-266-1-15. URL: "
        "https://cyberleninka.ru/article/n/osobennosti-proektirovaniya-mikroservisno-sobytiynyh-arhitektur-dlya-vysokonagruzhennyh-raspredelennyh-sistem-obrabotki-informatsiy "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "3) Ключев А. О., Кустарев П. В., Платунов А. Е. Распределенные информационно-управляющие системы: "
        "учебное пособие [Электронный ресурс]. СПб.: Университет ИТМО, 2015. 58 с. URL: "
        "https://books.ifmo.ru/file/pdf/1724.pdf "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "4) Подгорная С. В., Калинина Е. С., Манохина Т. В., Ступаков С. А. Перспективы применения "
        "облачных технологий для хранения информации [Электронный ресурс] // Международный журнал "
        "гуманитарных и естественных наук. 2024. № 4-3 (91). URL: "
        "https://sciup.org/perspektivy-primenenija-oblachnyh-tehnologij-dlja-hranenija-informacii-170204944 "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "5) Хомоненко А. Д., Абу Хасан Р. О надежности и доступности объектных хранилищ данных "
        "[Электронный ресурс] // Интеллектуальные технологии на транспорте. 2023. № S1. С. 123-128. URL: "
        "https://cyberleninka.ru/article/n/o-nadezhnosti-i-dostupnosti-obektnyh-hranilisch-dannyh "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "6) Ястребов Л. Д. Анализ современных стандартов сжатия изображений [Электронный ресурс] // "
        "Фундаментальные и прикладные исследования: проблемы и результаты. 2013. URL: "
        "https://cyberleninka.ru/article/n/analiz-sovremennyh-standartov-szhatiya-izobrazheniy "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "7) Кузнецов А. В., Шишкина Э. Л. Цифровая обработка изображений: учебное пособие [Электронный ресурс]. "
        "Воронеж: Издательский дом ВГУ, 2023. 160 с. URL: "
        "https://www.researchgate.net/publication/367434055_Cifrovaa_obrabotka_izobrazenij "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "8) Голуб Ю. И., Старовойтов В. В. Оценка качества цифровых изображений [Электронный ресурс]. "
        "Минск: ОИПИ НАН Беларуси, 2023. 252 с. URL: "
        "https://www.researchgate.net/publication/376713718_Ocenka_kacestva_cifrovyh_izobrazenij "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "9) Буй Тхи Тху Чанг, Фан Нгок Хоанг, Спицын В. Г. Классификация изображений на основе применения "
        "цветовой информации, вейвлет-преобразования Хаара и многослойной нейронной сети [Электронный ресурс] "
        "// Проблемы информатики. 2011. № 2 (10). URL: "
        "https://cyberleninka.ru/article/n/klassifikatsiya-izobrazheniy-na-osnove-primeneniya-tsvetovoy-informatsii-veyvlet-preobrazovaniya-haara-i-mnogosloynoy-neyronnoy-seti "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "10) Зеленский А. А., Письменскова М. М., Воронин В. В. Алгоритм поиска изображений в виде "
        "хэш-функций на основе глубинных нейросетевых технологий [Электронный ресурс] // Доклады "
        "Томского государственного университета систем управления и радиоэлектроники. 2018. Т. 21, № 3. "
        "С. 57-62. DOI: 10.21293/1818-0442-2018-21-3-57-62. URL: "
        "https://journal.tusur.ru/ru/arhiv/3-2018/algoritm-poiska-izobrazheniy-v-vide-hesh-funktsiy-na-osnove-glubinnyh-neyrosetevyh-tehnologiy "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "11) Коршунова К. П. Задачи и методы автоматического описания изображений [Электронный ресурс] // "
        "Системы управления, связи и безопасности. 2018. № 1. С. 30-77. URL: "
        "http://sccs.intelgr.com/archive/2018-01/02-Korshunova.pdf "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "12) Saaty T. L. How to Make a Decision: The Analytic Hierarchy Process [Электронный ресурс] // "
        "Interfaces. 1994. Vol. 24, № 6. P. 19-43. DOI: 10.1287/inte.24.6.19. URL: "
        "https://pubsonline.informs.org/doi/abs/10.1287/inte.24.6.19 "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "13) cwebp | WebP [Электронный ресурс] // Google for Developers. URL: "
        "https://developers.google.com/speed/webp/docs/cwebp "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "14) AV1 Image File Format (AVIF) v1.2.0 [Электронный ресурс] // Alliance for Open Media. URL: "
        "https://aomedia.org/specifications/avif/ "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "15) Sandler M., Howard A., Zhu M., Zhmoginov A., Chen L.-C. MobileNetV2: Inverted Residuals "
        "and Linear Bottlenecks [Электронный ресурс] // Proceedings of the IEEE/CVF Conference on "
        "Computer Vision and Pattern Recognition. 2018. P. 4510-4520. URL: "
        "https://openaccess.thecvf.com/content_cvpr_2018/html/Sandler_MobileNetV2_Inverted_Residuals_CVPR_2018_paper.html "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "16) Portable Network Graphics (PNG) Specification (Third Edition) [Электронный ресурс] // W3C Recommendation. "
        "2025. URL: https://www.w3.org/TR/png-3/ "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "17) Wallace G. K. The JPEG Still Picture Compression Standard [Электронный ресурс] // Communications of the ACM. "
        "1991. Vol. 34, № 4. P. 30-44. DOI: 10.1145/103085.103089. URL: "
        "https://cacm.acm.org/research/the-jpeg-still-picture-compression-standard/ "
        f"(дата обращения: {ACCESS_DATE})."
    ),
]


CITATIONS = {
    48: [4],
    119: [6, 7],
    125: [13, 14],
    248: [4],
    253: [8, 12],
    254: [1, 11],
    279: [1, 2, 3],
    394: [10, 11],
    398: [5],
    465: [7, 13, 14, 16, 17],
    466: [13, 14, 16, 17],
    467: [6, 8, 13, 14, 16, 17],
    474: [8, 12],
    484: [12],
    821: [9, 15],
}

MATH_TEXT_TAG = "{http://schemas.openxmlformats.org/officeDocument/2006/math}t"


def find_last_text_run(paragraph):
    for run in reversed(paragraph.runs):
        if run.text:
            return run
    return None


def strip_trailing_punctuation(paragraph):
    run = find_last_text_run(paragraph)
    if run is None:
        return ""

    match = re.search(r"([.!?])\s*$", run.text)
    if not match:
        return ""

    punctuation = match.group(1)
    run.text = re.sub(r"[.!?]\s*$", "", run.text)
    return punctuation


def add_bookmark(paragraph, bookmark_name: str, bookmark_id: int):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), bookmark_name)

    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))

    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, anchor: str, text: str):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    r_style = OxmlElement("w:rStyle")
    r_style.set(qn("w:val"), "Hyperlink")
    r_pr.append(r_style)

    run.append(r_pr)

    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def delete_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    parent.remove(p)
    paragraph._p = paragraph._element = None


def append_citations(paragraph, numbers):
    punctuation = strip_trailing_punctuation(paragraph)
    paragraph.add_run(" ")

    for idx, number in enumerate(numbers):
        add_internal_hyperlink(paragraph, f"src{number:02d}", f"[{number}]")
        if idx < len(numbers) - 1:
            paragraph.add_run(", ")

    if punctuation:
        paragraph.add_run(punctuation)


def replace_source_list(document: Document):
    heading_idx = None
    for idx, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip() == "Список использованных источников":
            heading_idx = idx
            break

    if heading_idx is None:
        raise RuntimeError("Не найден раздел 'Список использованных источников'")

    trailing_paragraphs = list(document.paragraphs[heading_idx + 1 :])
    if len(trailing_paragraphs) < 10:
        raise RuntimeError("Ожидалось не менее 10 абзацев после заголовка списка источников")

    reusable_paragraphs = trailing_paragraphs

    for source_text, paragraph in zip(SOURCES[: len(reusable_paragraphs)], reusable_paragraphs):
        paragraph.text = source_text

    if len(SOURCES) > len(reusable_paragraphs):
        for source_text in SOURCES[len(reusable_paragraphs) :]:
            document.add_paragraph(source_text)

    all_source_paragraphs = []
    for paragraph in document.paragraphs[heading_idx + 1 :]:
        if paragraph.text.strip():
            all_source_paragraphs.append(paragraph)

    if len(all_source_paragraphs) != len(SOURCES):
        raise RuntimeError("Не удалось сформировать итоговый список из 15 источников")

    for paragraph in list(document.paragraphs[heading_idx + 1 :]):
        if not paragraph.text.strip():
            delete_paragraph(paragraph)

    all_source_paragraphs = []
    for paragraph in document.paragraphs[heading_idx + 1 :]:
        if paragraph.text.strip():
            all_source_paragraphs.append(paragraph)

    bookmark_id = 5000
    for number, paragraph in enumerate(all_source_paragraphs, start=1):
        add_bookmark(paragraph, f"src{number:02d}", bookmark_id)
        bookmark_id += 1


def apply_citations(document: Document):
    for idx, numbers in CITATIONS.items():
        append_citations(document.paragraphs[idx], numbers)


def replace_math_text(paragraph, old: str, new: str):
    for element in paragraph._p.iter():
        if element.tag == MATH_TEXT_TAG and element.text == old:
            element.text = new
            return
    raise RuntimeError(f"Не удалось найти математический фрагмент '{old}'")


def update_chapter_two(document: Document):
    document.paragraphs[30].text = (
        "Тема выпускной квалификационной работы «Облачный сервис для архивного хранения фотографий». "
        "Отчет состоит из 88 страниц, 17 источников, 20 таблиц, 50 рисунков."
    )
    document.paragraphs[465].text = (
        "На основе анализа свойств изображений различных типов, а также характеристик форматов JPEG, PNG, WebP и AVIF "
        "в работе принято следующее распределение предпочтительных форматов сжатия. Такое распределение не является "
        "готовой табличной нормой из одного источника, а представляет собой проектное решение, сформированное по "
        "сравнительным характеристикам указанных кодеков"
    )
    document.paragraphs[466].text = (
        "Для фотографических изображений в качестве основных рассматриваются форматы WebP и AVIF, поскольку они "
        "предназначены для эффективного сжатия полноцветных изображений с сохранением визуального качества. Для "
        "текстовой графики и снимков интерфейса приоритет отдаётся PNG как формату без потерь с поддержкой "
        "прозрачности; в качестве дополнительного варианта рассматривается WebP, для которого официальный инструмент "
        "cwebp предусматривает специальные режимы text и drawing. Для иллюстраций используются WebP, AVIF и PNG, "
        "так как для этого класса изображений важно сочетать компактность хранения, поддержку прозрачности и "
        "возможность выбора между lossy- и lossless-представлением. Для изображений смешанного типа в качестве "
        "универсальных вариантов рассматриваются WebP и JPEG как наиболее распространённые форматы сжатия с потерями."
    )
    document.paragraphs[467].text = (
        "Для каждого формата задаётся дискретный набор значений качества. В качестве опорных уровней для lossy-"
        "кодеков используются типовые практические значения около 75 как базового уровня сжатия и диапазон 90-100 "
        "как уровня повышенного визуального качества. В документации WebP значение q=75 указано как стандартное "
        "значение по умолчанию, а q=100 соответствует наилучшему качеству; для JPEG уровень качества 75 "
        "рассматривается в сравнительных исследованиях как обычно рекомендуемый, а 95 характеризует режим высокого "
        "битрейта. Для PNG значение 100 в модели используется как условный маркер режима без потерь. Для AVIF "
        "диапазон качеств сдвинут немного ниже по сравнению с WebP, поскольку этот формат обеспечивает меньший "
        "размер файла при сопоставимом визуальном восприятии. Конкретные промежуточные значения 72, 78, 84, 90, 96 "
        "и аналогичные им множества для других форматов образуют редкую сетку перебора вокруг опорных уровней 75, "
        "85, 90 и 95, что позволяет уменьшить вычислительную сложность подбора без полного перебора всех значений от "
        "0 до 100"
    )
    document.paragraphs[474].text = (
        "После этого для каждого результата вычисляется оценка визуального качества и оценка степени уменьшения "
        "размера. Визуальная составляющая определяется через среднеквадратическую ошибку, преобразованную в "
        "нормированную метрику качества, как формула 2.15"
    )
    document.paragraphs[479].text = (
        "где S(·) - функция размера файла. В программной реализации значение Z(x, y_f,q) дополнительно "
        "ограничивается интервалом [0; 1], чтобы исключить отрицательные значения и значения больше единицы."
    )
    document.paragraphs[484].text = (
        "где α определяет вклад визуального качества изображения, а β - вклад степени уменьшения размера файла. "
        "В рассматриваемой реализации принята нормированная пара коэффициентов α = 0,9 и β = 0,1. Такое "
        "соотношение соответствует приоритету визуального качества над дополнительным уменьшением объёма файла и "
        "может быть интерпретировано как результат многокритериального выбора по методу анализа иерархий при "
        "отношении важности критериев 9:1."
    )
    document.paragraphs[535].text = (
        "После выбора рекомендуемых параметров система выполняет преобразование исходного изображения в целевой "
        "формат. В подразделах 2.2.3-2.2.6 приводятся математические модели используемых форматов JPEG, WebP, PNG "
        "и AVIF. В программной реализации непосредственное кодирование выполняется готовыми энкодерами библиотеки "
        "Imagick и внешней утилитой avifenc, поэтому данные формулы описывают внутреннюю логику применяемых "
        "кодеков, а не ручную реализацию всех их стадий в PHP-коде."
    )
    replace_math_text(document.paragraphs[483], "α=0,92,    β=0,08", "α=0,9,    β=0,1")


def set_run_font(run_element):
    r_pr = run_element.find(qn("w:rPr"))
    if r_pr is None:
        r_pr = OxmlElement("w:rPr")
        run_element.insert(0, r_pr)

    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)

    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        r_fonts.set(qn(attr), "Times New Roman")

    sz = r_pr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        r_pr.append(sz)
    sz.set(qn("w:val"), "28")

    sz_cs = r_pr.find(qn("w:szCs"))
    if sz_cs is None:
        sz_cs = OxmlElement("w:szCs")
        r_pr.append(sz_cs)
    sz_cs.set(qn("w:val"), "28")


def normalize_document_font(document: Document):
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(14)
    normal_style._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal_style._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(14)
        for run_element in paragraph._p.iter(qn("w:r")):
            set_run_font(run_element)


def main():
    src = Path(r"C:\Users\aricr\Desktop\институт 4\8 семестр\ВКР Калмыкова А.М. гр 220621.docx")
    dst = Path(r"C:\Users\aricr\Desktop\Microservices\vkr_code_aligned.docx")

    document = Document(src)
    update_chapter_two(document)
    apply_citations(document)
    replace_source_list(document)
    normalize_document_font(document)
    document.save(dst)
    print(dst)


if __name__ == "__main__":
    main()
