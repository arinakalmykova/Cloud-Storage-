from __future__ import annotations

from pathlib import Path

from docx import Document

import update_vkr_sources as base


ACCESS_DATE = base.ACCESS_DATE
MATH_TEXT_TAG = base.MATH_TEXT_TAG


SOURCES = [
    (
        "1) Ворсин В. А. Микро-сервисная архитектура бизнес-приложений: перспективы и проблемы "
        "[Электронный ресурс] // GLOBUS. 2020. С. 51-53. URL: "
        "https://cyberleninka.ru/article/n/mikro-servisnaya-arhitektura-biznes-prilozheniy-perspektivy-i-problemy "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "2) Хомоненко А. Д., Абу Хасан Р. О надежности и доступности объектных хранилищ данных "
        "[Электронный ресурс] // Интеллектуальные технологии на транспорте. 2023. № S1. С. 123-128. URL: "
        "https://cyberleninka.ru/article/n/o-nadezhnosti-i-dostupnosti-obektnyh-hranilisch-dannyh "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "3) Ястребов Л. Д. Анализ современных стандартов сжатия изображений [Электронный ресурс] // "
        "Фундаментальные и прикладные исследования: проблемы и результаты. 2013. URL: "
        "https://cyberleninka.ru/article/n/analiz-sovremennyh-standartov-szhatiya-izobrazheniy "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "4) Кузнецов А. В., Шишкина Э. Л. Цифровая обработка изображений: учебное пособие [Электронный ресурс]. "
        "Воронеж: Издательский дом ВГУ, 2023. 160 с. URL: "
        "https://www.researchgate.net/publication/367434055_Cifrovaa_obrabotka_izobrazenij "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "5) Голуб Ю. И., Старовойтов В. В. Оценка качества цифровых изображений [Электронный ресурс]. "
        "Минск: ОИПИ НАН Беларуси, 2023. 252 с. URL: "
        "https://www.researchgate.net/publication/376713718_Ocenka_kacestva_cifrovyh_izobrazenij "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "6) Коршунова К. П. Задачи и методы автоматического описания изображений [Электронный ресурс] // "
        "Системы управления, связи и безопасности. 2018. № 1. С. 30-77. URL: "
        "http://sccs.intelgr.com/archive/2018-01/02-Korshunova.pdf "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "7) Saaty T. L. How to Make a Decision: The Analytic Hierarchy Process [Электронный ресурс] // "
        "Interfaces. 1994. Vol. 24, № 6. P. 19-43. DOI: 10.1287/inte.24.6.19. URL: "
        "https://pubsonline.informs.org/doi/abs/10.1287/inte.24.6.19 "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "8) cwebp | WebP [Электронный ресурс] // Google for Developers. URL: "
        "https://developers.google.com/speed/webp/docs/cwebp "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "9) AV1 Image File Format (AVIF) v1.2.0 [Электронный ресурс] // Alliance for Open Media. URL: "
        "https://aomedia.org/specifications/avif/ "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "10) Sandler M., Howard A., Zhu M., Zhmoginov A., Chen L.-C. MobileNetV2: Inverted Residuals "
        "and Linear Bottlenecks [Электронный ресурс] // Proceedings of the IEEE/CVF Conference on "
        "Computer Vision and Pattern Recognition. 2018. P. 4510-4520. URL: "
        "https://openaccess.thecvf.com/content_cvpr_2018/html/Sandler_MobileNetV2_Inverted_Residuals_CVPR_2018_paper.html "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "11) Portable Network Graphics (PNG) Specification (Third Edition) [Электронный ресурс] // "
        "W3C Recommendation. 2025. URL: https://www.w3.org/TR/png-3/ "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "12) Wallace G. K. The JPEG Still Picture Compression Standard [Электронный ресурс] // "
        "Communications of the ACM. 1991. Vol. 34, № 4. P. 30-44. DOI: 10.1145/103085.103089. URL: "
        "https://cacm.acm.org/research/the-jpeg-still-picture-compression-standard/ "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "13) WebP Compression Study [Электронный ресурс] // Google for Developers. URL: "
        "https://developers.google.com/speed/webp/docs/webp_study "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "14) Using AVIF to compress images on your site [Электронный ресурс] // web.dev. URL: "
        "https://web.dev/articles/compress-images-avif?hl=en "
        f"(дата обращения: {ACCESS_DATE})."
    ),
    (
        "15) Wang Z., Bovik A. C., Sheikh H. R., Simoncelli E. P. Image Quality Assessment: "
        "From Error Visibility to Structural Similarity [Электронный ресурс] // IEEE Transactions on "
        "Image Processing. 2004. Vol. 13, № 4. P. 600-612. DOI: 10.1109/TIP.2003.819861. URL: "
        "https://pubmed.ncbi.nlm.nih.gov/15376593/ "
        f"(дата обращения: {ACCESS_DATE})."
    ),
]


CITATIONS = {
    119: [3, 4],
    125: [3, 8, 9, 11, 12],
    253: [5, 15],
    254: [4, 6],
    279: [1],
    394: [6],
    398: [2],
    465: [10],
    477: [8, 9, 11, 12, 14],
    478: [8, 9, 11, 14],
    479: [8, 9, 12, 13, 14],
    486: [5, 15],
    496: [7],
    505: [13, 14, 15],
}


def set_math_texts(paragraph, values: list[str]) -> None:
    math_nodes = [element for element in paragraph._p.iter() if element.tag == MATH_TEXT_TAG]
    if len(math_nodes) != len(values):
        raise RuntimeError(
            f"Для абзаца ожидалось {len(math_nodes)} математических токенов, получено {len(values)}"
        )
    for node, value in zip(math_nodes, values):
        node.text = value


def set_formula_line(paragraph, formula_text: str, number: str) -> None:
    paragraph.text = f"\t{formula_text}\t({number})"


def apply_citations(document: Document) -> None:
    for idx, numbers in CITATIONS.items():
        base.append_citations(document.paragraphs[idx], numbers)


def update_chapter_two(document: Document) -> None:
    document.paragraphs[30].text = (
        "Тема выпускной квалификационной работы «Облачный сервис для архивного хранения фотографий». "
        "Отчет состоит из 90 страниц, 15 источников, 20 таблиц, 50 рисунков."
    )

    document.paragraphs[468].text = (
        "После получения класса система определяет рекомендуемый формат хранения и диапазон качества сжатия."
    )
    document.paragraphs[471].text = (
        "В программной реализации рекомендуемый формат задается отображением типа изображения в основной кодек "
        "(формула 2.11)."
    )
    document.paragraphs[472].text = "\tfrec(photo) = AVIF,"
    document.paragraphs[473].text = "\tfrec(text_graphics) = PNG,"
    document.paragraphs[474].text = "\tfrec(ui_screenshot) = PNG\t(2.11)"
    document.paragraphs[475].text = "\tfrec(illustration) = WebP,"
    document.paragraphs[476].text = "\tfrec(mixed) = WebP."

    document.paragraphs[477].text = (
        "Выбор рекомендуемого формата выполнен на основе анализа свойств изображений различных типов и "
        "характеристик современных форматов JPEG, PNG, WebP и AVIF. Такое сопоставление не является готовой "
        "табличной нормой из одного источника, а представляет собой проектное правило, выведенное из "
        "сравнительных свойств кодеков и зафиксированное в программной реализации сервиса."
    )
    document.paragraphs[478].text = (
        "Для фотографических изображений в качестве основного формата выбран AVIF, поскольку этот формат "
        "ориентирован на высокую эффективность сжатия полноцветных изображений при сохранении визуального "
        "качества. Для текстовой графики и снимков интерфейса выбран PNG как формат без потерь, хорошо "
        "подходящий для текста, резких границ и однородных областей цвета. Для иллюстраций выбран WebP, "
        "поскольку он поддерживает как lossy-, так и lossless-представление, прозрачность и специализированные "
        "режимы кодирования drawing и text. Для изображений смешанного типа WebP используется как универсальный "
        "формат по умолчанию."
    )
    document.paragraphs[479].text = (
        "Для каждого формата задается диапазон quality и шаг первичного перебора (формула 2.12). Эти диапазоны "
        "не являются универсальными нормативными константами стандартов, а задают инженерные интервалы поиска, "
        "сформированные вокруг документированных и практически используемых уровней качества: для JPEG опорным "
        "служит уровень около 75, для WebP используется параметр q с типовым значением по умолчанию 75, для "
        "AVIF диапазон поиска смещается в область 60-90, а для PNG значение 100 используется как условный "
        "маркер режима без потерь. В программной реализации грубый перебор выполняется по редкой сетке с шагом "
        "5, что позволяет ограничить время поиска без полного перебора всех значений от 0 до 100."
    )

    document.paragraphs[486].text = (
        "После этого для каждого результата вычисляются оценка визуального качества и оценка степени уменьшения "
        "размера. В текущей реализации визуальная составляющая определяется через индекс структурного сходства "
        "SSIM, который лучше согласуется с человеческим восприятием качества, чем метрики на основе ошибки."
    )
    set_formula_line(document.paragraphs[487], "V(x, yf,q) = SSIM(x, yf,q)", "2.15")
    document.paragraphs[488].text = (
        "где SSIM(x, yf,q) - индекс структурного сходства между исходным и сжатым изображением."
    )
    document.paragraphs[491].text = (
        "где S(·) - функция размера файла. В программной реализации значение Z(x, yf,q) дополнительно "
        "ограничивается интервалом [0; 1], чтобы исключить отрицательные значения и значения больше единицы."
    )

    document.paragraphs[494].text = (
        "Для интегральной оценки кандидата в сервисе также используется взвешенная формула 2.18."
    )
    set_formula_line(document.paragraphs[495], "α = 0,9,    β = 0,1", "2.18")
    document.paragraphs[496].text = (
        "где α определяет вклад визуального качества изображения, а β - вклад степени уменьшения размера файла. "
        "В рассматриваемой реализации принята нормированная пара коэффициентов α = 0,9 и β = 0,1. Такое "
        "соотношение соответствует приоритету визуального качества над дополнительным уменьшением объема файла и "
        "может быть интерпретировано как результат многокритериального выбора по методу анализа иерархий при "
        "отношении важности критериев 9:1."
    )

    document.paragraphs[497].text = (
        "На первом этапе из результатов грубого перебора формируется множество допустимых кандидатов, для которых "
        "SSIM не ниже 0,95. Внутри этого множества выбирается sweet spot, то есть момент, когда дальнейший рост "
        "quality почти не улучшает визуальное качество, но уже заметно увеличивает размер файла."
    )
    set_formula_line(document.paragraphs[498], "Ωa(x) = {(f, q) ∈ Ω(x) | V(x, yf,q) ≥ 0,95}", "2.19")
    document.paragraphs[499].text = (
        "После этого система выполняет дополнительное уточнение качества в окрестности найденного значения q*. "
        "Для выбранного формата f* строится локальное множество кандидатов с радиусом 4 и шагом 2."
    )
    set_formula_line(
        document.paragraphs[500],
        "Ωl(x) = {(f*, q) | q ∈ {q* - 4, q* - 2, q*, q* + 2, q* + 4} ∩ [0; 100]}",
        "2.20",
    )
    document.paragraphs[501].text = (
        "Далее для локального множества снова выполняются пробное сжатие, вычисление SSIM и проверка sweet spot."
    )
    set_formula_line(document.paragraphs[502], "q' = sweetspot(Ωl(x))", "2.21")
    document.paragraphs[503].text = (
        "Итоговая рекомендация определяется как выбранный формат f* и уточненное значение качества q'."
    )
    set_formula_line(document.paragraphs[504], "R(x) = (f*, q')", "2.22")
    document.paragraphs[505].text = (
        "В итоге алгоритм рекомендации включает последовательность из грубого перебора по диапазону quality, "
        "фильтрации кандидатов по порогу SSIM >= 0,95, выбора sweet spot и локального уточнения качества. Это "
        "позволяет находить параметр сжатия, при котором визуальное качество остается приемлемым, а дальнейшее "
        "увеличение quality уже не дает соразмерного выигрыша."
    )

    document.paragraphs[547].text = (
        "После выбора рекомендуемых параметров система выполняет преобразование исходного изображения в целевой "
        "формат. В подразделах 2.2.3-2.2.6 приводятся математические модели форматов JPEG, WebP, PNG и AVIF. "
        "В программной реализации непосредственное кодирование выполняется готовыми энкодерами библиотеки "
        "Imagick и внешней утилитой avifenc, поэтому приведенные формулы описывают внутреннюю логику используемых "
        "кодеков, а не ручную реализацию всех стадий кодирования в PHP-коде."
    )


def build_document(src: Path, dst: Path) -> None:
    document = Document(src)
    base.SOURCES = SOURCES
    update_chapter_two(document)
    apply_citations(document)
    base.replace_source_list(document)
    base.normalize_document_font(document)
    document.save(dst)


def main() -> None:
    src = Path(r"C:\Users\aricr\Desktop\Microservices\source_vkr.docx")
    dst = Path(r"C:\Users\aricr\Desktop\Microservices\vkr_report_final.docx")
    build_document(src, dst)
    print(dst)


if __name__ == "__main__":
    main()
